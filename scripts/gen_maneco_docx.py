#!/usr/bin/env python3.12
"""Generate DOCX for Introduction to Managerial Economics."""
import json, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE     = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SLUG     = "introduction-to-managerial-economics"
JSON_IN  = os.path.join(BASE, f"content/subjects/managerial-economics/lectures/{SLUG}.json")
DOCX_OUT = os.path.join(BASE, f"public/downloads/{SLUG}.docx")

with open(JSON_IN) as f:
    d = json.load(f)

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

def shade(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def h1(doc, text, color="1E3A5F"):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h

def h2(doc, text, color="1D4ED8"):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h

def body(doc, text, bg=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(6)
    if bg:
        shade(p, bg)
    return p

# ── Cover ─────────────────────────────────────────────────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run(d["title"])
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_paragraph()

mp = doc.add_paragraph()
mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
mp.add_run(
    f"Subject: Managerial Economics  |  Difficulty: {d['difficulty'].title()}  "
    f"|  Read time: {d['estimatedReadTime']}  |  Date: {d['date']}"
).font.size = Pt(9)

doc.add_paragraph()

# ── 1. Summary ────────────────────────────────────────────────────────────────
h1(doc, "1. Lecture Summary")
body(doc, d["summary"], "EFF6FF")

# ── 2. Key Concepts ───────────────────────────────────────────────────────────
h1(doc, "2. Key Concepts")
for c in d["concepts"]:
    li = doc.add_paragraph(style="List Bullet")
    li.add_run(c).font.size = Pt(10.5)

# ── 3. Detailed Notes ─────────────────────────────────────────────────────────
h1(doc, "3. Detailed Notes")
for i, s in enumerate(d["notes"], 1):
    h2(doc, f"{i}. {s['heading']}")
    body(doc, s["content"])

# ── 4. Real-World Example ─────────────────────────────────────────────────────
h1(doc, "4. Real-World Example")
body(doc, d["example"], "F0FDF4")

# ── 5. Case Study ─────────────────────────────────────────────────────────────
h1(doc, "5. Case Study")
body(doc, d["caseStudy"], "FFFBEB")

# ── 6. Key Takeaways ──────────────────────────────────────────────────────────
h1(doc, "6. Key Takeaways")
for i, t in enumerate(d["takeaways"], 1):
    li = doc.add_paragraph(style="List Number")
    li.add_run(t).font.size = Pt(10.5)

# ── 7. Knowledge Quiz ─────────────────────────────────────────────────────────
doc.add_page_break()
h1(doc, f"7. Knowledge Quiz  ({len(d['quiz'])} Questions)")
labels = ["A", "B", "C", "D"]
for i, q in enumerate(d["quiz"], 1):
    qp = doc.add_paragraph()
    qp.paragraph_format.space_before = Pt(10)
    qr = qp.add_run(f"Q{i}. {q['question']}")
    qr.bold = True
    qr.font.size = Pt(10.5)
    for j, opt in enumerate(q["options"]):
        op = doc.add_paragraph()
        op.paragraph_format.left_indent = Inches(0.3)
        op.paragraph_format.space_after = Pt(2)
        marker = "✓ " if j == q["correct"] else "    "
        or_ = op.add_run(f"{marker}{labels[j]}. {opt}")
        or_.font.size = Pt(10)
        if j == q["correct"]:
            or_.bold = True
            or_.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
    ep = doc.add_paragraph()
    ep.paragraph_format.left_indent = Inches(0.3)
    ep.paragraph_format.space_after = Pt(6)
    er = ep.add_run(f"Explanation: {q['explanation']}")
    er.italic = True
    er.font.size = Pt(9.5)
    er.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.save(DOCX_OUT)
print(f"DOCX saved → {DOCX_OUT}")
