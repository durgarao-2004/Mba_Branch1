#!/usr/bin/env python3.12
"""Generate DOCX lecture file from JSON content."""
import json, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
JSON_PATH = os.path.join(BASE, "content/subjects/financial-management/lectures/introduction-to-financial-management.json")
OUT_PATH  = os.path.join(BASE, "public/downloads/introduction-to-financial-management.docx")

with open(JSON_PATH, "r") as f:
    d = json.load(f)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper: set paragraph shading ────────────────────────────────────────────
def shade_para(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def add_section_heading(doc, text, level=1, color_hex="1E3A5F"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return h

def add_body(doc, text, shade=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(6)
    if shade:
        shade_para(p, shade)
    return p

# ── Cover / Title block ───────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run(d["title"])
title_run.bold = True
title_run.font.size = Pt(22)
title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_paragraph()

meta_para = doc.add_paragraph()
meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_text = (
    f"Subject: Financial Management  |  Difficulty: {d['difficulty'].title()}  "
    f"|  Read time: {d['estimatedReadTime']}  |  Date: {d['date']}"
)
meta_run = meta_para.add_run(meta_text)
meta_run.font.size = Pt(9)
meta_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

tags_para = doc.add_paragraph()
tags_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
tags_run = tags_para.add_run("Tags: " + "  •  ".join(d.get("tags", [])))
tags_run.font.size = Pt(9)
tags_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_paragraph()

# ── 1. Lecture Summary ────────────────────────────────────────────────────────
add_section_heading(doc, "1. Lecture Summary", 1)
p = doc.add_paragraph()
p.add_run(d["summary"]).font.size = Pt(10.5)
p.paragraph_format.space_after = Pt(8)
shade_para(p, "EFF6FF")

# ── 2. Key Concepts ───────────────────────────────────────────────────────────
add_section_heading(doc, "2. Key Concepts", 1)
for concept in d["concepts"]:
    item = doc.add_paragraph(style="List Bullet")
    run = item.add_run(concept)
    run.font.size = Pt(10.5)

# ── 3. Detailed Notes ─────────────────────────────────────────────────────────
add_section_heading(doc, "3. Detailed Notes", 1)
for i, section in enumerate(d["notes"], 1):
    h = doc.add_heading(f"{i}. {section['heading']}", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    p = doc.add_paragraph()
    p.add_run(section["content"]).font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(8)

# ── 4. Real-World Example ─────────────────────────────────────────────────────
add_section_heading(doc, "4. Real-World Example", 1)
p = doc.add_paragraph()
p.add_run(d["example"]).font.size = Pt(10.5)
shade_para(p, "F0FDF4")

# ── 5. Case Study ─────────────────────────────────────────────────────────────
add_section_heading(doc, "5. Case Study", 1)
p = doc.add_paragraph()
p.add_run(d["caseStudy"]).font.size = Pt(10.5)
shade_para(p, "FFFBEB")

# ── 6. Key Takeaways ──────────────────────────────────────────────────────────
add_section_heading(doc, "6. Key Takeaways", 1)
for i, takeaway in enumerate(d["takeaways"], 1):
    item = doc.add_paragraph(style="List Number")
    run = item.add_run(takeaway)
    run.font.size = Pt(10.5)

# ── 7. Knowledge Quiz ─────────────────────────────────────────────────────────
doc.add_page_break()
add_section_heading(doc, "7. Knowledge Quiz", 1)
intro = doc.add_paragraph()
intro.add_run(f"{len(d['quiz'])} MCQs — Test your understanding of {d['title']}").font.size = Pt(10)
intro.runs[0].italic = True

for i, q in enumerate(d["quiz"], 1):
    # Question
    qp = doc.add_paragraph()
    qp.paragraph_format.space_before = Pt(10)
    qr = qp.add_run(f"Q{i}. {q['question']}")
    qr.bold = True
    qr.font.size = Pt(10.5)

    # Options
    labels = ["A", "B", "C", "D"]
    for j, option in enumerate(q["options"]):
        op = doc.add_paragraph()
        op.paragraph_format.left_indent = Inches(0.3)
        op.paragraph_format.space_after = Pt(2)
        marker = "✓ " if j == q["correct"] else "    "
        or_ = op.add_run(f"{marker}{labels[j]}. {option}")
        or_.font.size = Pt(10)
        if j == q["correct"]:
            or_.bold = True
            or_.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

    # Explanation
    ep = doc.add_paragraph()
    ep.paragraph_format.left_indent = Inches(0.3)
    ep.paragraph_format.space_after = Pt(6)
    er = ep.add_run(f"Explanation: {q['explanation']}")
    er.italic = True
    er.font.size = Pt(9.5)
    er.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.save(OUT_PATH)
print(f"DOCX saved → {OUT_PATH}")
